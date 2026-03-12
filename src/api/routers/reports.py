"""보고서 관련 라우터"""
import json
import logging
import os
import tempfile
import io
import base64
import time
import glob
import re
from datetime import datetime
from typing import Optional, List
from pathlib import Path
from fastapi import APIRouter, HTTPException, Query
from fastapi.responses import FileResponse
from pydantic import BaseModel
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION
from PIL import Image
import requests
from moviepy.video.io.VideoFileClip import VideoFileClip
from config.settings import CLIPS_DIR, VIDEOS_DIR, CONVERTED_VIDEOS_DIR, REPORTS_DIR
from database.connection import get_db_connection, verify_user_exists
from dependencies import verify_user_dependency
from exceptions import NotFoundException, ValidationException, DatabaseException
from fastapi import Depends

logger = logging.getLogger(__name__)

# PDF 변환 라이브러리 import (Windows: docx2pdf, Linux: LibreOffice)
try:
    from docx2pdf import convert
    PDF_CONVERSION_AVAILABLE = True
    PDF_CONVERSION_METHOD = "docx2pdf"
except ImportError:
    PDF_CONVERSION_AVAILABLE = False
    PDF_CONVERSION_METHOD = None
    logger.warning("docx2pdf가 설치되지 않았습니다. PDF 변환 기능을 사용할 수 없습니다.")

def add_horizontal_line(paragraph):
    """단락에 구분선(horizontal rule) 추가 (Word에서 --- 입력 후 엔터와 동일한 효과)"""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    
    # 단락 테두리 요소 생성
    pBdr = OxmlElement('w:pBdr')
    
    # 다른 속성들보다 앞에 삽입
    pPr.insert_element_before(
        pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    
    # 하단 테두리 요소 생성 (구분선)
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')  # 단일 선 스타일
    bottom.set(qn('w:sz'), '6')        # 선 두께 (1/8 pt 단위, 6 = 0.75pt)
    bottom.set(qn('w:space'), '1')     # 텍스트와 테두리 사이 간격
    bottom.set(qn('w:color'), 'auto')  # 테두리 색상 (자동)
    
    # 하단 테두리를 단락 테두리에 추가
    pBdr.append(bottom)
    
    return paragraph

router = APIRouter()

# ==================== 요청 모델 ====================
class CreateReportRequest(BaseModel):
    user_id: str
    title: str
    description: Optional[str] = None
    content: str
    word_count: int = 0
    video_ids: Optional[List[int]] = None
    video_titles: Optional[List[str]] = None

class CreateReportResponse(BaseModel):
    success: bool
    report_id: Optional[int] = None
    message: str

class ClipData(BaseModel):
    id: Optional[str] = None
    title: str
    url: str
    sentence: Optional[str] = None
    start_time: Optional[float] = None
    end_time: Optional[float] = None
    sourceVideo: Optional[str] = None

class CreateWordReportRequest(BaseModel):
    user_id: str
    title: str
    author: Optional[str] = None
    description: Optional[str] = None
    query: Optional[str] = None
    clips: List[ClipData]

class AddClipsToReportRequest(BaseModel):
    user_id: str
    clips: List[ClipData]

class UpdateReportRequest(BaseModel):
    user_id: str
    title: Optional[str] = None
    description: Optional[str] = None
    content: Optional[str] = None

# ==================== 엔드포인트 ====================
@router.get("/check-title")
async def check_report_title(
    title: str = Query(..., description="확인할 보고서 제목"),
    user_id: str = Depends(verify_user_dependency)
):
    """보고서 제목 중복 확인"""
    try:
        if not title:
            raise ValidationException("보고서 제목이 필요합니다.")
        
        with get_db_connection() as cursor:
            # 동일한 제목의 보고서가 있는지 확인
            cursor.execute("""
                SELECT COUNT(*) FROM vss_reports
                WHERE USER_ID = ? AND TITLE = ?
            """, (user_id, title))
            
            count = cursor.fetchone()[0]
        
        return {
            "success": True,
            "exists": count > 0,
            "message": "이미 존재하는 보고서 제목입니다." if count > 0 else "사용 가능한 제목입니다."
        }
    except (ValidationException, NotFoundException):
        raise
    except Exception as e:
        logger.error(f"보고서 제목 확인 실패: {e}")
        raise DatabaseException(f"보고서 제목 확인 중 오류가 발생했습니다: {str(e)}")

@router.post("", response_model=CreateReportResponse)
async def create_report(request: CreateReportRequest):
    """보고서 생성"""
    try:
        user_id = request.user_id
        title = request.title
        description = request.description or ""
        content = request.content
        word_count = request.word_count or 0
        video_ids = request.video_ids or []
        video_titles = request.video_titles or []
        
        if not user_id:
            raise HTTPException(status_code=400, detail="사용자 ID가 필요합니다.")
        if not title:
            raise HTTPException(status_code=400, detail="보고서 제목이 필요합니다.")
        if not content:
            raise HTTPException(status_code=400, detail="보고서 내용이 필요합니다.")
        
        verify_user_exists(user_id)
        
        # 보고서 저장
        video_ids_json = json.dumps(video_ids) if video_ids else None
        video_titles_json = json.dumps(video_titles) if video_titles else None
        
        with get_db_connection() as cursor:
            cursor.execute("""
                INSERT INTO vss_reports (USER_ID, TITLE, DESCRIPTION, CONTENT, WORD_COUNT, VIDEO_IDS, VIDEO_TITLES)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            """, (user_id, title, description, content, word_count, video_ids_json, video_titles_json))
            
            report_id = cursor.lastrowid
        
        logger.info(f"보고서 생성 완료: USER_ID={user_id}, REPORT_ID={report_id}, TITLE={title}")
        
        return {
            "success": True,
            "report_id": report_id,
            "message": "보고서가 성공적으로 생성되었습니다."
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"보고서 생성 실패: {e}")
        raise HTTPException(status_code=500, detail=f"보고서 생성 중 오류가 발생했습니다: {str(e)}")

@router.post("/create-word")
async def create_word_report(request: CreateWordReportRequest):
    """워드 파일 형식의 보고서 생성"""
    try:
        user_id = request.user_id
        title = request.title
        author = request.author or user_id  # 작성자가 없으면 user_id 사용
        description = request.description or ""
        query = request.query or ""
        clips = request.clips
        
        if not user_id:
            raise HTTPException(status_code=400, detail="사용자 ID가 필요합니다.")
        if not title:
            raise HTTPException(status_code=400, detail="보고서 제목이 필요합니다.")
        if not clips or len(clips) == 0:
            raise HTTPException(status_code=400, detail="클립 데이터가 필요합니다.")
        
        verify_user_exists(user_id)
        
        # 워드 문서 생성
        doc = Document()
        
        # 페이지 여백 설정 (상하좌우 1인치)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # ==================== 보고서 구조 ====================
        # 1. 보고서 제목
        title_heading = doc.add_heading(title, level=1)
        title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_heading.runs[0]
        title_run.font.size = Pt(24)
        title_run.bold = True
        title_run.font.color.rgb = RGBColor(0, 51, 102)  # 진한 파란색
        
        doc.add_paragraph()  # 빈 줄
        doc.add_paragraph()  # 빈 줄
        
        # 구분선 추가
        separator_para = doc.add_paragraph()
        add_horizontal_line(separator_para)
        
        doc.add_paragraph()  # 빈 줄
        doc.add_paragraph()  # 빈 줄
        
        # 2. 작성자 및 작성 일자 (테이블 형태)
        current_date = datetime.now().strftime("%Y년 %m월 %d일")
        info_table = doc.add_table(rows=2, cols=2)
        info_table.style = 'Light Grid Accent 1'
        
        # 작성자 행
        author_row = info_table.rows[0]
        author_row.cells[0].text = "작성자"
        author_row.cells[0].paragraphs[0].runs[0].bold = True
        author_row.cells[0].paragraphs[0].runs[0].font.size = Pt(12)
        author_row.cells[1].text = author
        author_row.cells[1].paragraphs[0].runs[0].font.size = Pt(12)
        
        # 작성 일자 행
        date_row = info_table.rows[1]
        date_row.cells[0].text = "작성 일자"
        date_row.cells[0].paragraphs[0].runs[0].bold = True
        date_row.cells[0].paragraphs[0].runs[0].font.size = Pt(12)
        date_row.cells[1].text = current_date
        date_row.cells[1].paragraphs[0].runs[0].font.size = Pt(12)
        
        doc.add_paragraph()  # 빈 줄
        doc.add_paragraph()  # 빈 줄
        
        # 구분선 추가
        separator_para2 = doc.add_paragraph()
        add_horizontal_line(separator_para2)
        
        doc.add_paragraph()  # 빈 줄
        doc.add_paragraph()  # 빈 줄
        
        # 3. Q. 검색어 섹션
        query_section_para = doc.add_paragraph()
        query_label = query_section_para.add_run("Q. 검색어")
        query_label.bold = True
        query_label.font.size = Pt(16)
        query_label.font.color.rgb = RGBColor(0, 102, 204)  # 파란색
        
        doc.add_paragraph()  # 빈 줄
        
        # 검색어 내용 박스 (들여쓰기)
        query_content_para = doc.add_paragraph()
        query_content_para.paragraph_format.left_indent = Inches(0.5)
        query_content_para.paragraph_format.space_before = Pt(6)
        query_content_para.paragraph_format.space_after = Pt(6)
        
        query_content = query if query else "검색어가 제공되지 않았습니다."
        query_content_run = query_content_para.add_run(query_content)
        query_content_run.font.size = Pt(13)
        
        doc.add_paragraph()  # 빈 줄
        doc.add_paragraph()  # 빈 줄
        
        # 구분선 추가
        separator_para3 = doc.add_paragraph()
        add_horizontal_line(separator_para3)
        
        doc.add_paragraph()  # 빈 줄
        doc.add_paragraph()  # 빈 줄
        
        # 4. A. 검색 결과 섹션
        answer_para = doc.add_paragraph()
        answer_label = answer_para.add_run("A. 검색 결과")
        answer_label.bold = True
        answer_label.font.size = Pt(16)
        answer_label.font.color.rgb = RGBColor(0, 102, 204)  # 파란색
        
        doc.add_paragraph()  # 빈 줄
        
        # 검색 결과 클립 및 설명
        for idx, clip in enumerate(clips, 1):
            # 클립 항목 구분선 (첫 번째 항목이 아닌 경우)
            if idx > 1:
                clip_separator = doc.add_paragraph()
                clip_separator.paragraph_format.space_before = Pt(12)
                clip_separator.paragraph_format.space_after = Pt(12)
                add_horizontal_line(clip_separator)
                doc.add_paragraph()  # 빈 줄
            
            # 클립 번호 및 설명 (들여쓰기)
            clip_para = doc.add_paragraph()
            clip_para.paragraph_format.left_indent = Inches(0.3)
            clip_para.paragraph_format.space_before = Pt(8)
            clip_para.paragraph_format.space_after = Pt(4)
            
            clip_number = clip_para.add_run(f"[{idx}] ")
            clip_number.bold = True
            clip_number.font.size = Pt(13)
            clip_number.font.color.rgb = RGBColor(51, 51, 51)  # 진한 회색
            
            # 클립 설명 텍스트
            sentence = clip.sentence if clip.sentence else clip.title
            if sentence:
                clip_desc = clip_para.add_run(sentence)
                clip_desc.font.size = Pt(13)
            
            # 시간 정보가 있으면 추가
            if clip.start_time is not None and clip.end_time is not None:
                time_info = f" (시간: {clip.start_time:.2f}초 - {clip.end_time:.2f}초)"
                time_run = clip_para.add_run(time_info)
                time_run.font.size = Pt(11)
                time_run.italic = True
                time_run.font.color.rgb = RGBColor(128, 128, 128)  # 회색
            
            doc.add_paragraph()  # 빈 줄
            
            # 썸네일 이미지 추가 시도 (들여쓰기)
            if clip.url:
                try:
                    image_data = None
                    if clip.url.lower().endswith(('.jpg', '.jpeg', '.png', '.gif', '.webp')):
                        image_data = download_image(clip.url)
                    else:
                        thumbnail_time = clip.start_time if clip.start_time is not None else 0.0
                        image_data = get_video_thumbnail(clip.url, thumbnail_time)
                    
                    if image_data:
                        try:
                            image = Image.open(io.BytesIO(image_data))
                            max_width = 5.0  # 인치 단위 (약간 더 크게)
                            width, height = image.size
                            # 너비가 최대 너비보다 크면 리사이즈
                            if width > max_width * 72:  # 72 DPI 기준
                                ratio = (max_width * 72) / width
                                new_width = int(width * ratio)
                                new_height = int(height * ratio)
                                image = image.resize((new_width, new_height), Image.Resampling.LANCZOS)
                            
                            img_byte_arr = io.BytesIO()
                            image.save(img_byte_arr, format='PNG')
                            img_byte_arr.seek(0)
                            
                            # 이미지 단락에 들여쓰기 추가
                            img_para = doc.add_paragraph()
                            img_para.paragraph_format.left_indent = Inches(0.5)
                            img_para.paragraph_format.space_before = Pt(6)
                            img_para.paragraph_format.space_after = Pt(12)
                            
                            # 이미지를 단락에 추가
                            run = img_para.add_run()
                            run.add_picture(img_byte_arr, width=Inches(max_width))
                        except Exception as e:
                            logger.warning(f"이미지 추가 실패 (클립 {idx}): {e}")
                except Exception as e:
                    logger.warning(f"썸네일 처리 실패 (클립 {idx}, URL: {clip.url}): {e}")
            
            doc.add_paragraph()  # 빈 줄
        
        # 보고서 내용을 텍스트로 변환 (데이터베이스 저장용) - 새로운 구조
        report_content = f"# {title}\n\n"
        report_content += f"**작성자:** {author}\n\n"
        report_content += f"**작성 일자:** {current_date}\n\n"
        report_content += "=" * 50 + "\n\n"
        report_content += "## Q. 검색어\n\n"
        report_content += f"{query if query else '검색어가 제공되지 않았습니다.'}\n\n"
        report_content += "=" * 50 + "\n\n"
        report_content += "## A. 검색 결과\n\n"
        
        for idx, clip in enumerate(clips, 1):
            sentence = clip.sentence if clip.sentence else clip.title
            if sentence:
                report_content += f"{idx}. {sentence}"
                if clip.start_time is not None and clip.end_time is not None:
                    report_content += f" ({clip.start_time:.2f}초 - {clip.end_time:.2f}초)"
                report_content += "\n\n"
        
        # 단어 수 계산
        word_count = len(report_content.split())
        
        # 데이터베이스에 보고서 저장
        video_ids = []
        video_titles = []
        for clip in clips:
            if clip.id:
                try:
                    video_ids.append(int(clip.id))
                except (ValueError, TypeError):
                    pass
            if clip.title:
                video_titles.append(clip.title)
        
        video_ids_json = json.dumps(video_ids) if video_ids else None
        video_titles_json = json.dumps(video_titles) if video_titles else None
        
        # Word 파일 저장
        REPORTS_DIR.mkdir(exist_ok=True, parents=True)
        
        # 파일명 생성 (보고서 ID를 사용하기 위해 먼저 DB에 저장)
        # 임시로 타임스탬프 사용
        timestamp = int(time.time() * 1000)
        safe_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).rstrip()
        safe_title = safe_title.replace(' ', '_')[:50]  # 파일명 길이 제한
        filename = f"{safe_title}_{timestamp}.docx"
        file_path = REPORTS_DIR / filename
        
        # Word 문서 저장
        doc.save(str(file_path))
        
        # 파일 URL 생성 (정적 파일 서빙용)
        file_url = f"/reports-files/{filename}"
        
        # 데이터베이스에 보고서 저장 (Word 파일 경로 포함)
        with get_db_connection() as cursor:
            cursor.execute("""
                INSERT INTO vss_reports (USER_ID, TITLE, DESCRIPTION, CONTENT, WORD_COUNT, VIDEO_IDS, VIDEO_TITLES)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            """, (user_id, title, description, report_content, word_count, video_ids_json, video_titles_json))
            
            report_id = cursor.lastrowid
        
        # 보고서 ID를 사용하여 파일명 재생성 및 파일명 변경
        final_filename = f"report_{report_id}_{timestamp}.docx"
        final_file_path = REPORTS_DIR / final_filename
        if file_path.exists():
            file_path.rename(final_file_path)
        final_file_url = f"/reports-files/{final_filename}"
        
        logger.info(f"워드 보고서 생성 완료: USER_ID={user_id}, REPORT_ID={report_id}, TITLE={title}, CLIPS={len(clips)}, FILE={final_filename}")
        
        return {
            "success": True,
            "report_id": report_id,
            "file_url": final_file_url,
            "message": "보고서가 성공적으로 생성되었습니다."
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"워드 보고서 생성 실패: {e}")
        raise HTTPException(status_code=500, detail=f"보고서 생성 중 오류가 발생했습니다: {str(e)}")

@router.get("")
async def get_reports(
    user_id: str = Query(..., description="사용자 ID"),
    page: int = Query(1, ge=1, description="페이지 번호"),
    page_size: int = Query(10, ge=1, le=100, description="페이지당 항목 수")
):
    """보고서 목록 조회 (페이지네이션 지원)"""
    try:
        if not user_id:
            raise HTTPException(status_code=400, detail="사용자 ID가 필요합니다.")
        
        verify_user_exists(user_id)
        
        with get_db_connection() as cursor:
            # vss_reports 테이블이 없으면 빈 목록 반환
            try:
                cursor.execute("SELECT COUNT(*) FROM vss_reports WHERE USER_ID = ?", (user_id,))
            except Exception as e:
                # 테이블이 없으면 빈 목록 반환
                logger.warning(f"vss_reports 테이블이 없습니다: {e}")
                return {
                    "success": True,
                    "reports": [],
                    "total": 0,
                    "page": page,
                    "page_size": page_size,
                    "pages": 0
                }
            
            # 전체 개수 조회
            cursor.execute("SELECT COUNT(*) FROM vss_reports WHERE USER_ID = ?", (user_id,))
            total = cursor.fetchone()[0]
            
            # 페이지네이션 계산
            offset = (page - 1) * page_size
            pages = max(1, (total + page_size - 1) // page_size)
            
            # 보고서 목록 조회 (최신순)
            cursor.execute("""
                SELECT ID, TITLE, DESCRIPTION, CONTENT, WORD_COUNT, VIDEO_IDS, VIDEO_TITLES, CREATED_AT, UPDATED_AT
                FROM vss_reports
                WHERE USER_ID = ?
                ORDER BY CREATED_AT DESC
                LIMIT ? OFFSET ?
            """, (user_id, page_size, offset))
            
            rows = cursor.fetchall()
        
        # context manager 밖에서 데이터 처리
        reports = []
        for row in rows:
            report_id, title, description, content, word_count, video_ids_json, video_titles_json, created_at, updated_at = row
            
            # Word 파일 경로 생성 (보고서 ID 기반)
            # 실제 파일 찾기
            pattern = str(REPORTS_DIR / f"report_{report_id}_*.docx")
            matching_files = glob.glob(pattern)
            if matching_files:
                filename = os.path.basename(matching_files[0])
                file_url = f"/reports-files/{filename}"
            else:
                file_url = None
            
            # JSON 문자열 파싱
            video_ids = json.loads(video_ids_json) if video_ids_json else []
            video_titles = json.loads(video_titles_json) if video_titles_json else []
            
            reports.append({
                "id": report_id,
                "report_id": report_id,
                "title": title,
                "description": description or "",
                "content": content or "",
                "word_count": word_count or 0,
                "video_ids": video_ids,
                "video_titles": video_titles,
                "file_url": file_url,
                "created_at": created_at.isoformat() if created_at else None,
                "createdAt": created_at.isoformat() if created_at else None,
                "updated_at": updated_at.isoformat() if updated_at else None
            })
        
        logger.info(f"보고서 목록 조회 완료: USER_ID={user_id}, 총 {total}개, 페이지 {page}/{pages}")
        
        return {
            "success": True,
            "reports": reports,
            "total": total,
            "page": page,
            "page_size": page_size,
            "pages": pages
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"보고서 목록 조회 실패: {e}")
        raise HTTPException(status_code=500, detail=f"보고서 목록 조회 중 오류가 발생했습니다: {str(e)}")


@router.get("/{report_id}")
async def get_report(
    report_id: int,
    user_id: str = Query(..., description="사용자 ID")
):
    """보고서 상세 조회"""
    try:
        if not user_id:
            raise HTTPException(status_code=400, detail="사용자 ID가 필요합니다.")
        
        verify_user_exists(user_id)
        
        with get_db_connection() as cursor:
            # 보고서 조회
            cursor.execute("""
                SELECT ID, TITLE, DESCRIPTION, CONTENT, WORD_COUNT, VIDEO_IDS, VIDEO_TITLES, CREATED_AT, UPDATED_AT
                FROM vss_reports
                WHERE ID = ? AND USER_ID = ?
            """, (report_id, user_id))
            
            row = cursor.fetchone()
        
        if not row:
            raise HTTPException(status_code=404, detail="보고서를 찾을 수 없습니다.")
        
        report_id_db, title, description, content, word_count, video_ids_json, video_titles_json, created_at, updated_at = row
        
        # JSON 문자열 파싱
        video_ids = json.loads(video_ids_json) if video_ids_json else []
        video_titles = json.loads(video_titles_json) if video_titles_json else []
        
        # Word 파일 경로 찾기
        pattern = str(REPORTS_DIR / f"report_{report_id_db}_*.docx")
        matching_files = glob.glob(pattern)
        if matching_files:
            filename = os.path.basename(matching_files[0])
            file_url = f"/reports-files/{filename}"
        else:
            file_url = None
        
        return {
            "success": True,
            "report": {
                "id": report_id_db,
                "report_id": report_id_db,
                "title": title,
                "description": description or "",
                "content": content or "",
                "word_count": word_count or 0,
                "video_ids": video_ids,
                "video_titles": video_titles,
                "file_url": file_url,
                "created_at": created_at.isoformat() if created_at else None,
                "createdAt": created_at.isoformat() if created_at else None,
                "updated_at": updated_at.isoformat() if updated_at else None
            }
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"보고서 상세 조회 실패: {e}")
        raise HTTPException(status_code=500, detail=f"보고서 상세 조회 중 오류가 발생했습니다: {str(e)}")

@router.delete("/{report_id}")
async def delete_report(
    report_id: int,
    user_id: str = Query(..., description="사용자 ID")
):
    """보고서 삭제"""
    try:
        if not user_id:
            raise HTTPException(status_code=400, detail="사용자 ID가 필요합니다.")
        
        verify_user_exists(user_id)
        
        with get_db_connection() as cursor:
            # 먼저 보고서 존재 여부 및 소유권 확인
            cursor.execute("""
                SELECT ID FROM vss_reports
                WHERE ID = ? AND USER_ID = ?
            """, (report_id, user_id))
            
            existing_report = cursor.fetchone()
            
            if not existing_report:
                raise HTTPException(status_code=404, detail="보고서를 찾을 수 없거나 권한이 없습니다.")
        
        # Word 파일 삭제 시도
        pattern = str(REPORTS_DIR / f"report_{report_id}_*.docx")
        matching_files = glob.glob(pattern)
        for file_path in matching_files:
            try:
                if os.path.exists(file_path):
                    os.unlink(file_path)
                    logger.info(f"Word 파일 삭제 완료: {file_path}")
            except Exception as e:
                logger.warning(f"Word 파일 삭제 실패 ({file_path}): {e}")
        
        # 보고서 삭제
        with get_db_connection() as cursor:
            cursor.execute("""
                DELETE FROM vss_reports
                WHERE ID = ? AND USER_ID = ?
            """, (report_id, user_id))
            
            deleted_count = cursor.rowcount
        
        if deleted_count == 0:
            # 이미 삭제되었을 수 있음
            logger.warning(f"보고서 삭제 실패: USER_ID={user_id}, REPORT_ID={report_id} (이미 삭제되었을 수 있음)")
            return {
                "success": True,
                "message": "보고서가 이미 삭제되었거나 존재하지 않습니다."
            }
        
        logger.info(f"보고서 삭제 완료: USER_ID={user_id}, REPORT_ID={report_id}")
        
        return {
            "success": True,
            "message": "보고서가 성공적으로 삭제되었습니다."
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"보고서 삭제 실패: {e}")
        raise HTTPException(status_code=500, detail=f"보고서 삭제 중 오류가 발생했습니다: {str(e)}")

@router.put("/{report_id}/add-clips")
async def add_clips_to_report(
    report_id: int,
    request: AddClipsToReportRequest
):
    """기존 보고서에 클립 추가"""
    try:
        user_id = request.user_id
        clips = request.clips
        
        if not user_id:
            raise HTTPException(status_code=400, detail="사용자 ID가 필요합니다.")
        if not clips or len(clips) == 0:
            raise HTTPException(status_code=400, detail="클립 데이터가 필요합니다.")
        
        verify_user_exists(user_id)
        
        with get_db_connection() as cursor:
            # 기존 보고서 확인
            cursor.execute("""
                SELECT ID, TITLE, DESCRIPTION, CONTENT, WORD_COUNT, VIDEO_IDS, VIDEO_TITLES
                FROM vss_reports
                WHERE ID = ? AND USER_ID = ?
            """, (report_id, user_id))
            
            row = cursor.fetchone()
            if not row:
                raise HTTPException(status_code=404, detail="보고서를 찾을 수 없거나 권한이 없습니다.")
            
            report_id_db, title, description, existing_content, word_count, video_ids_json, video_titles_json = row
        
        # 기존 클립 개수 계산 (콘텐츠에서 ## 숫자. 패턴으로 찾기)
        existing_clip_count = len(re.findall(r'## \d+\.', existing_content))
        
        # 기존 클립 ID와 URL 추출 (중복 체크용)
        existing_video_ids = json.loads(video_ids_json) if video_ids_json else []
        existing_clip_urls = set()
        
        # 기존 콘텐츠에서 클립 URL 추출 (마크다운 형식에서)
        url_pattern = r'- 비디오 URL: (https?://[^\s]+)'
        existing_urls = re.findall(url_pattern, existing_content)
        existing_clip_urls.update(existing_urls)
        
        # 기존 콘텐츠에서 클립 제목도 추출 (제목으로도 중복 체크)
        existing_clip_titles = set()
        title_pattern = r'## \d+\. ([^\n]+)'
        existing_titles = re.findall(title_pattern, existing_content)
        existing_clip_titles.update([t.strip() for t in existing_titles])
        
        # 중복 클립 필터링 (URL 또는 ID 기준)
        new_clips = []
        duplicate_clips = []
        
        for clip in clips:
            is_duplicate = False
            
            # URL로 중복 체크
            if clip.url and clip.url in existing_clip_urls:
                is_duplicate = True
            
            # ID로 중복 체크
            if clip.id:
                try:
                    clip_id = int(clip.id)
                    if clip_id in existing_video_ids:
                        is_duplicate = True
                except (ValueError, TypeError):
                    pass
            
            # 제목으로도 중복 체크 (URL과 ID가 없는 경우)
            if not is_duplicate and clip.title and clip.title.strip() in existing_clip_titles:
                is_duplicate = True
            
            if is_duplicate:
                duplicate_clips.append(clip)
            else:
                new_clips.append(clip)
        
        # 중복 클립이 모두인 경우
        if len(new_clips) == 0:
            return {
                "success": False,
                "message": f"추가할 수 있는 새로운 클립이 없습니다. 모든 클립이 이미 보고서에 포함되어 있습니다.",
                "duplicate_count": len(duplicate_clips)
            }
        
        # 일부만 중복인 경우 경고 메시지 포함
        duplicate_message = ""
        if len(duplicate_clips) > 0:
            duplicate_message = f" ({len(duplicate_clips)}개의 중복 클립은 제외되었습니다)"
        
        # 새 클립만 사용
        clips = new_clips
        
        # 기존 Word 파일 찾기
        pattern = str(REPORTS_DIR / f"report_{report_id}_*.docx")
        matching_files = glob.glob(pattern)
        if not matching_files:
            raise HTTPException(status_code=404, detail="Word 파일을 찾을 수 없습니다.")
        
        word_file_path = matching_files[0]
        
        # 기존 Word 문서 열기
        doc = Document(str(word_file_path))
        
        # 기존 클립 수 확인 (제목 개수로 추정)
        existing_clip_count = len([p for p in doc.paragraphs if p.style.name == 'Heading 1' and p.text.strip() and not p.text.strip().startswith('#')])
        
        # 새로운 클립 추가
        clip_images = []
        new_clip_start_idx = existing_clip_count + 1
        
        for idx, clip in enumerate(clips, new_clip_start_idx):
            # 구분선 추가 (이전 클립과 구분)
            separator_para = doc.add_paragraph()
            add_horizontal_line(separator_para)
            
            # 클립 번호 및 제목
            clip_heading = doc.add_heading(f"{idx}. {clip.title}", level=1)
            
            # 시간 정보
            if clip.start_time is not None and clip.end_time is not None:
                time_para = doc.add_paragraph(f"시간: {format_time(clip.start_time)} - {format_time(clip.end_time)}")
                time_para.style = 'List Bullet'
            
            # 소스 비디오 정보
            if clip.sourceVideo:
                source_para = doc.add_paragraph(f"소스: {clip.sourceVideo}")
                source_para.style = 'List Bullet'
            
            # 썸네일 이미지 추가 시도
            image_data = None
            if clip.url:
                try:
                    if clip.url.lower().endswith(('.jpg', '.jpeg', '.png', '.gif', '.webp')):
                        image_data = download_image(clip.url)
                    else:
                        thumbnail_time = clip.start_time if clip.start_time is not None else 0.0
                        image_data = get_video_thumbnail(clip.url, thumbnail_time)
                    
                    if image_data:
                        try:
                            image = Image.open(io.BytesIO(image_data))
                            max_width = 4.0
                            width, height = image.size
                            if width > max_width * 72:
                                ratio = (max_width * 72) / width
                                new_width = int(width * ratio)
                                new_height = int(height * ratio)
                                image = image.resize((new_width, new_height), Image.Resampling.LANCZOS)
                            
                            img_byte_arr = io.BytesIO()
                            image.save(img_byte_arr, format='PNG')
                            img_byte_arr.seek(0)
                            
                            doc.add_picture(img_byte_arr, width=Inches(max_width))
                            
                            img_byte_arr.seek(0)
                            clip_images.append(base64.b64encode(img_byte_arr.read()).decode('utf-8'))
                        except Exception as e:
                            logger.warning(f"이미지 추가 실패: {e}")
                            clip_images.append(None)
                    else:
                        clip_images.append(None)
                except Exception as e:
                    logger.warning(f"썸네일 처리 실패 ({clip.url}): {e}")
                    clip_images.append(None)
            else:
                clip_images.append(None)
            
            # 장면 설명 추가
            if clip.sentence:
                sentence_para = doc.add_paragraph("장면 설명:")
                sentence_para.style = 'Heading 3'
                desc_para = doc.add_paragraph(clip.sentence)
                desc_para.style = 'Normal'
        
        # 기존 내용에 새 클립 내용 추가
        new_content = existing_content
        for idx, clip in enumerate(clips, new_clip_start_idx):
            new_content += f"\n## {idx}. {clip.title}\n\n"
            if clip.start_time is not None and clip.end_time is not None:
                new_content += f"- 시간: {format_time(clip.start_time)} - {format_time(clip.end_time)}\n"
            if clip.sourceVideo:
                new_content += f"- 소스: {clip.sourceVideo}\n"
            
            clip_image_idx = idx - new_clip_start_idx
            if clip_image_idx < len(clip_images) and clip_images[clip_image_idx]:
                new_content += f"\n![썸네일 {idx}](data:image/png;base64,{clip_images[clip_image_idx]})\n\n"
            elif clip.url:
                new_content += f"- 비디오 URL: {clip.url}\n"
            
            if clip.sentence:
                new_content += f"\n### 장면 설명:\n{clip.sentence}\n\n"
            new_content += "=" * 50 + "\n\n"
        
        # 단어 수 재계산
        word_count = len(new_content.split())
        
        # 기존 VIDEO_IDS와 VIDEO_TITLES에 새 클립 추가
        existing_video_ids = json.loads(video_ids_json) if video_ids_json else []
        existing_video_titles = json.loads(video_titles_json) if video_titles_json else []
        
        for clip in clips:
            if clip.id:
                try:
                    clip_id = int(clip.id)
                    if clip_id not in existing_video_ids:
                        existing_video_ids.append(clip_id)
                except (ValueError, TypeError):
                    pass
            if clip.title and clip.title not in existing_video_titles:
                existing_video_titles.append(clip.title)
        
        video_ids_json = json.dumps(existing_video_ids) if existing_video_ids else None
        video_titles_json = json.dumps(existing_video_titles) if existing_video_titles else None
        
        # description 업데이트 (클립 개수 반영)
        # 기존 클립 수와 새로 추가된 클립 수를 합산
        total_clip_count = existing_clip_count + len(clips)
        
        # description이 "n개의 클립 검색 결과" 형식인지 확인
        description_pattern_ko = r'(\d+)개의 클립 검색 결과'
        description_pattern_en = r'Search results for (\d+) clips?'
        
        updated_description = description
        if description:
            # 한국어 패턴 매칭
            match_ko = re.search(description_pattern_ko, description)
            if match_ko:
                updated_description = re.sub(description_pattern_ko, f'{total_clip_count}개의 클립 검색 결과', description)
            else:
                # 영어 패턴 매칭
                match_en = re.search(description_pattern_en, description)
                if match_en:
                    updated_description = re.sub(description_pattern_en, f'Search results for {total_clip_count} clips', description)
                else:
                    # 패턴이 없으면 description에 클립 개수 정보 추가
                    # 한국어/영어 판단
                    is_korean = any(ord(c) >= 0xAC00 and ord(c) <= 0xD7A3 for c in description) if description else True
                    if is_korean:
                        updated_description = f"{total_clip_count}개의 클립 검색 결과"
                    else:
                        updated_description = f"Search results for {total_clip_count} clips"
        else:
            # description이 없으면 새로 생성 (기본값은 한국어)
            updated_description = f"{total_clip_count}개의 클립 검색 결과"
        
        # Word 파일의 description도 업데이트
        # Word 문서에서 description 부분 찾기 및 업데이트
        if doc.paragraphs:
            # 제목 다음 단락들을 확인 (description은 보통 제목 바로 다음)
            for i, para in enumerate(doc.paragraphs):
                para_text = para.text.strip()
                # description 패턴이 있는 단락 찾기
                if re.search(description_pattern_ko, para_text) or re.search(description_pattern_en, para_text):
                    # 단락 텍스트 업데이트
                    new_text = re.sub(description_pattern_ko, f'{total_clip_count}개의 클립 검색 결과', para_text)
                    new_text = re.sub(description_pattern_en, f'Search results for {total_clip_count} clips', new_text)
                    if new_text != para_text:
                        para.clear()
                        para.add_run(new_text)
                    break
                # 제목 다음 3개 단락 내에서만 확인 (성능 최적화)
                if i > 5:  # 제목, 작성자, 구분선 등을 고려하여 적절한 범위
                    break
        
        # Word 파일 저장
        doc.save(str(word_file_path))
        
        # 데이터베이스 업데이트 (description 포함)
        with get_db_connection() as cursor:
            cursor.execute("""
                UPDATE vss_reports
                SET CONTENT = ?, DESCRIPTION = ?, WORD_COUNT = ?, VIDEO_IDS = ?, VIDEO_TITLES = ?, UPDATED_AT = CURRENT_TIMESTAMP
                WHERE ID = ? AND USER_ID = ?
            """, (new_content, updated_description, word_count, video_ids_json, video_titles_json, report_id, user_id))
        
        logger.info(f"보고서에 클립 추가 완료: USER_ID={user_id}, REPORT_ID={report_id}, 추가된 클립 수={len(clips)}")
        
        return {
            "success": True,
            "report_id": report_id,
            "file_url": f"/reports-files/{os.path.basename(word_file_path)}",
            "message": f"{len(clips)}개의 클립이 보고서에 추가되었습니다.{duplicate_message}",
            "added_count": len(clips),
            "duplicate_count": len(duplicate_clips) if 'duplicate_clips' in locals() else 0
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"보고서에 클립 추가 실패: {e}")
        raise HTTPException(status_code=500, detail=f"보고서에 클립 추가 중 오류가 발생했습니다: {str(e)}")

@router.put("/{report_id}")
async def update_report(
    report_id: int,
    request: UpdateReportRequest
):
    """보고서 내용 수정"""
    try:
        user_id = request.user_id
        
        if not user_id:
            raise HTTPException(status_code=400, detail="사용자 ID가 필요합니다.")
        
        verify_user_exists(user_id)
        
        with get_db_connection() as cursor:
            # 기존 보고서 확인
            cursor.execute("""
                SELECT ID, TITLE, DESCRIPTION, CONTENT, WORD_COUNT
                FROM vss_reports
                WHERE ID = ? AND USER_ID = ?
            """, (report_id, user_id))
            
            row = cursor.fetchone()
            if not row:
                raise HTTPException(status_code=404, detail="보고서를 찾을 수 없거나 권한이 없습니다.")
            
            report_id_db, existing_title, existing_description, existing_content, existing_word_count = row
        
        # 업데이트할 필드 준비
        title = request.title if request.title is not None else existing_title
        description = request.description if request.description is not None else existing_description
        content = request.content if request.content is not None else existing_content
        
        # 단어 수 재계산 (내용이 변경된 경우)
        word_count = existing_word_count
        if request.content is not None:
            word_count = len(content.split())
        
        # 기존 Word 파일 경로 확인 (파일 시스템에서 찾기)
        pattern = str(REPORTS_DIR / f"report_{report_id}_*.docx")
        matching_files = glob.glob(pattern)
        existing_file_url = None
        if matching_files:
            filename = os.path.basename(matching_files[0])
            existing_file_url = f"/reports-files/{filename}"
        
        # Word 파일 재생성 (내용이 변경된 경우)
        new_file_url = existing_file_url
        if request.content is not None and content != existing_content:
            try:
                # 마크다운 내용을 Word 문서로 변환
                doc = Document()
            
                lines = content.split('\n')
                current_paragraph = None
                in_list = False
                
                for line in lines:
                    original_line = line
                    line = line.strip()
                    
                    # 빈 줄 처리
                    if not line:
                        if current_paragraph and not in_list:
                            current_paragraph = None
                        in_list = False
                        continue
                    
                    # 제목 처리 (# ## ###)
                    if line.startswith('#'):
                        in_list = False
                        level = len(line) - len(line.lstrip('#'))
                        text = line.lstrip('#').strip()
                        if text:
                            doc.add_heading(text, level=min(level, 3))
                            current_paragraph = None
                    # 이미지 처리 (![alt](data:image/...))
                    elif '![' in line and 'data:image' in line:
                        in_list = False
                        # base64 이미지 추출 (여러 줄에 걸쳐 있을 수 있음)
                        match = re.search(r'data:image/([^;]+);base64,([A-Za-z0-9+/=]+)', line)
                        if match:
                            img_type, img_data = match.groups()
                            try:
                                img_bytes = base64.b64decode(img_data)
                                img_io = io.BytesIO(img_bytes)
                                doc.add_picture(img_io, width=Inches(4))
                                current_paragraph = None
                            except Exception as e:
                                logger.warning(f"이미지 추가 실패: {e}")
                    # 리스트 항목 처리 (- * 숫자.)
                    elif re.match(r'^[-*]\s+', line) or re.match(r'^\d+\.\s+', line):
                        in_list = True
                        text = re.sub(r'^[-*]\s+', '', line)
                        text = re.sub(r'^\d+\.\s+', '', text)
                        if text:
                            para = doc.add_paragraph(text, style='List Bullet')
                            current_paragraph = para
                    # 구분선 처리 (=== 또는 ---)
                    elif re.match(r'^={3,}$', line) or re.match(r'^-{3,}$', line):
                        in_list = False
                        separator_para = doc.add_paragraph()
                        add_horizontal_line(separator_para)
                        current_paragraph = None
                    # 볼드/이탤릭 처리 (**text** 또는 *text*)
                    elif '**' in line or '*' in line:
                        in_list = False
                        # 간단한 볼드 처리
                        text = line.replace('**', '').replace('*', '')
                        if text:
                            para = doc.add_paragraph()
                            para.add_run(text)
                            current_paragraph = para
                    # 일반 텍스트
                    else:
                        if in_list:
                            in_list = False
                        if current_paragraph and not in_list:
                            current_paragraph.add_run(' ' + line)
                        else:
                            current_paragraph = doc.add_paragraph(line)
                
                # Word 파일 저장
                REPORTS_DIR.mkdir(exist_ok=True, parents=True)
                timestamp = int(time.time() * 1000)
                final_filename = f"report_{report_id}_{timestamp}.docx"
                final_file_path = REPORTS_DIR / final_filename
                doc.save(str(final_file_path))
                
                # 기존 파일 삭제 (있는 경우)
                if existing_file_url:
                    try:
                        old_filename = existing_file_url.replace("/reports-files/", "")
                        old_file_path = REPORTS_DIR / old_filename
                        if old_file_path.exists() and old_file_path != final_file_path:
                            old_file_path.unlink()
                            logger.info(f"기존 Word 파일 삭제: {old_filename}")
                    except Exception as e:
                        logger.warning(f"기존 Word 파일 삭제 실패: {e}")
                
                new_file_url = f"/reports-files/{final_filename}"
                logger.info(f"Word 파일 재생성 완료: {final_filename}")
            except Exception as e:
                logger.error(f"Word 파일 재생성 실패: {e}")
                # Word 파일 재생성 실패해도 DB 업데이트는 계속 진행
        
        # 데이터베이스 업데이트
        with get_db_connection() as cursor:
            cursor.execute("""
                UPDATE vss_reports
                SET TITLE = ?, DESCRIPTION = ?, CONTENT = ?, WORD_COUNT = ?, UPDATED_AT = CURRENT_TIMESTAMP
                WHERE ID = ? AND USER_ID = ?
            """, (title, description, content, word_count, report_id, user_id))
        
        logger.info(f"보고서 수정 완료: USER_ID={user_id}, REPORT_ID={report_id}")
        
        return {
            "success": True,
            "report_id": report_id,
            "file_url": new_file_url,
            "message": "보고서가 성공적으로 수정되었습니다."
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"보고서 수정 실패: {e}")
        raise HTTPException(status_code=500, detail=f"보고서 수정 중 오류가 발생했습니다: {str(e)}")

def format_time(seconds: Optional[float]) -> str:
    """초를 MM:SS 형식으로 변환"""
    if seconds is None:
        return ""
    minutes = int(seconds // 60)
    secs = int(seconds % 60)
    return f"{minutes:02d}:{secs:02d}"

def download_image(url: str, timeout: int = 10) -> Optional[bytes]:
    """이미지 URL에서 이미지 다운로드"""
    try:
        response = requests.get(url, timeout=timeout, stream=True)
        if response.status_code == 200:
            return response.content
        return None
    except Exception as e:
        logger.warning(f"이미지 다운로드 실패 ({url}): {e}")
        return None

def get_video_thumbnail(video_url: str, time_seconds: float = 0.0) -> Optional[bytes]:
    """비디오 URL에서 썸네일 추출 (지정된 시간의 프레임)"""
    video_path = None
    is_temp_file = False
    
    try:
        # URL을 로컬 파일 경로로 변환 시도
        local_path = None
        
        # /clips/ URL인 경우
        if '/clips/' in video_url:
            filename = os.path.basename(video_url.split('/clips/')[-1])
            local_path = CLIPS_DIR / filename
        # /video-files/ URL인 경우
        elif '/video-files/' in video_url:
            filename = os.path.basename(video_url.split('/video-files/')[-1])
            local_path = VIDEOS_DIR / filename
        # /converted-videos/ URL인 경우
        elif '/converted-videos/' in video_url:
            filename = os.path.basename(video_url.split('/converted-videos/')[-1])
            local_path = CONVERTED_VIDEOS_DIR / filename
        
        # 로컬 파일 경로가 존재하는 경우 사용
        if local_path and os.path.exists(local_path):
            video_path = str(local_path)
        # 직접 파일 경로인 경우
        elif os.path.exists(video_url):
            video_path = video_url
        else:
            # URL인 경우 임시 파일로 다운로드
            try:
                response = requests.get(video_url, timeout=30, stream=True)
                if response.status_code != 200:
                    return None
                
                # 임시 파일 생성
                with tempfile.NamedTemporaryFile(delete=False, suffix='.mp4') as tmp_file:
                    for chunk in response.iter_content(chunk_size=8192):
                        tmp_file.write(chunk)
                    video_path = tmp_file.name
                    is_temp_file = True
            except Exception as e:
                logger.warning(f"비디오 다운로드 실패 ({video_url}): {e}")
                return None
        
        if not video_path or not os.path.exists(video_path):
            return None
        
        # moviepy로 비디오 열기
        video = VideoFileClip(video_path)
        
        # 지정된 시간의 프레임 추출 (기본값: 첫 프레임)
        frame_time = min(time_seconds, video.duration - 0.1) if video.duration else 0.0
        frame = video.get_frame(frame_time)
        video.close()
        
        # PIL Image로 변환
        img = Image.fromarray(frame)
        
        # 이미지 크기 조정 (최대 너비 800px)
        max_width = 800
        if img.width > max_width:
            ratio = max_width / img.width
            new_size = (max_width, int(img.height * ratio))
            img = img.resize(new_size, Image.Resampling.LANCZOS)
        
        # BytesIO로 변환
        img_byte_arr = io.BytesIO()
        img.save(img_byte_arr, format='PNG')
        img_byte_arr.seek(0)
        
        return img_byte_arr.getvalue()
    except Exception as e:
        logger.warning(f"비디오 썸네일 추출 실패 ({video_url}): {e}")
        return None
    finally:
        # 임시 파일 정리 (URL에서 다운로드한 경우만)
        if is_temp_file and video_path and os.path.exists(video_path):
            try:
                os.unlink(video_path)
            except:
                pass

def convert_docx_to_pdf(docx_path: str, pdf_path: Optional[str] = None) -> Optional[str]:
    """DOCX 파일을 PDF로 변환
    
    Args:
        docx_path: DOCX 파일 경로
        pdf_path: 출력 PDF 파일 경로 (None이면 자동 생성)
    
    Returns:
        PDF 파일 경로 (성공 시), None (실패 시)
    """
    if not PDF_CONVERSION_AVAILABLE:
        logger.error("PDF 변환 라이브러리가 설치되지 않았습니다.")
        return None
    
    try:
        docx_path_obj = Path(docx_path)
        if not docx_path_obj.exists():
            logger.error(f"DOCX 파일을 찾을 수 없습니다: {docx_path}")
            return None
        
        # PDF 경로가 지정되지 않으면 자동 생성
        if pdf_path is None:
            pdf_path = str(docx_path_obj.with_suffix('.pdf'))
        else:
            pdf_path_obj = Path(pdf_path)
            pdf_path_obj.parent.mkdir(parents=True, exist_ok=True)
        
        # PDF 변환 실행
        if PDF_CONVERSION_METHOD == "docx2pdf":
            # Windows: docx2pdf 사용 (MS Word 필요)
            convert(docx_path, pdf_path)
        else:
            logger.error(f"지원하지 않는 PDF 변환 방법: {PDF_CONVERSION_METHOD}")
            return None
        
        # 변환된 PDF 파일 확인
        if os.path.exists(pdf_path):
            logger.info(f"PDF 변환 완료: {docx_path} -> {pdf_path}")
            return pdf_path
        else:
            logger.error(f"PDF 변환 실패: 파일이 생성되지 않았습니다. {pdf_path}")
            return None
            
    except Exception as e:
        logger.error(f"PDF 변환 중 오류 발생: {e}")
        return None

@router.get("/{report_id}/pdf")
async def get_report_pdf(
    report_id: int,
    user_id: str = Query(..., description="사용자 ID")
):
    """보고서를 PDF로 변환하여 반환"""
    try:
        if not user_id:
            raise HTTPException(status_code=400, detail="사용자 ID가 필요합니다.")
        
        if not PDF_CONVERSION_AVAILABLE:
            raise HTTPException(status_code=503, detail="PDF 변환 기능을 사용할 수 없습니다. 서버에 docx2pdf가 설치되어 있는지 확인하세요.")
        
        verify_user_exists(user_id)
        
        with get_db_connection() as cursor:
            # 보고서 확인
            cursor.execute("""
                SELECT ID, TITLE, FILE_URL
                FROM vss_reports
                WHERE ID = ? AND USER_ID = ?
            """, (report_id, user_id))
            
            row = cursor.fetchone()
            if not row:
                raise HTTPException(status_code=404, detail="보고서를 찾을 수 없거나 권한이 없습니다.")
            
            report_id_db, title, file_url = row
        
        # Word 파일 경로 찾기
        pattern = str(REPORTS_DIR / f"report_{report_id_db}_*.docx")
        matching_files = glob.glob(pattern)
        if not matching_files:
            raise HTTPException(status_code=404, detail="Word 파일을 찾을 수 없습니다.")
        
        docx_file_path = matching_files[0]
        
        # PDF 파일 경로 생성
        pdf_file_path = str(Path(docx_file_path).with_suffix('.pdf'))
        
        # PDF 파일이 없거나 DOCX 파일이 더 최신인 경우 변환
        docx_mtime = os.path.getmtime(docx_file_path)
        pdf_exists = os.path.exists(pdf_file_path)
        pdf_mtime = os.path.getmtime(pdf_file_path) if pdf_exists else 0
        
        if not pdf_exists or docx_mtime > pdf_mtime:
            # PDF 변환 실행
            converted_pdf_path = convert_docx_to_pdf(docx_file_path, pdf_file_path)
            if not converted_pdf_path:
                raise HTTPException(status_code=500, detail="PDF 변환에 실패했습니다.")
        
        # PDF 파일 반환
        if not os.path.exists(pdf_file_path):
            raise HTTPException(status_code=404, detail="PDF 파일을 찾을 수 없습니다.")
        
        return FileResponse(
            pdf_file_path,
            media_type="application/pdf",
            filename=f"{title}_{report_id_db}.pdf"
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"PDF 변환 실패: {e}")
        raise HTTPException(status_code=500, detail=f"PDF 변환 중 오류가 발생했습니다: {str(e)}")

@router.post("/{report_id}/regenerate-pdf")
async def regenerate_report_pdf(
    report_id: int,
    user_id: str = Query(..., description="사용자 ID")
):
    """보고서 PDF를 강제로 재생성"""
    try:
        if not user_id:
            raise HTTPException(status_code=400, detail="사용자 ID가 필요합니다.")
        
        if not PDF_CONVERSION_AVAILABLE:
            raise HTTPException(status_code=503, detail="PDF 변환 기능을 사용할 수 없습니다.")
        
        verify_user_exists(user_id)
        
        with get_db_connection() as cursor:
            # 보고서 확인
            cursor.execute("""
                SELECT ID, TITLE
                FROM vss_reports
                WHERE ID = ? AND USER_ID = ?
            """, (report_id, user_id))
            
            row = cursor.fetchone()
            if not row:
                raise HTTPException(status_code=404, detail="보고서를 찾을 수 없거나 권한이 없습니다.")
            
            report_id_db, title = row
        
        # Word 파일 경로 찾기
        pattern = str(REPORTS_DIR / f"report_{report_id_db}_*.docx")
        matching_files = glob.glob(pattern)
        if not matching_files:
            raise HTTPException(status_code=404, detail="Word 파일을 찾을 수 없습니다.")
        
        docx_file_path = matching_files[0]
        pdf_file_path = str(Path(docx_file_path).with_suffix('.pdf'))
        
        # 기존 PDF 파일 삭제 (있는 경우)
        if os.path.exists(pdf_file_path):
            try:
                os.unlink(pdf_file_path)
            except Exception as e:
                logger.warning(f"기존 PDF 파일 삭제 실패: {e}")
        
        # PDF 변환 실행
        converted_pdf_path = convert_docx_to_pdf(docx_file_path, pdf_file_path)
        if not converted_pdf_path:
            raise HTTPException(status_code=500, detail="PDF 변환에 실패했습니다.")
        
        return {
            "success": True,
            "message": "PDF가 성공적으로 재생성되었습니다.",
            "pdf_url": f"/reports/{report_id}/pdf?user_id={user_id}"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"PDF 재생성 실패: {e}")
        raise HTTPException(status_code=500, detail=f"PDF 재생성 중 오류가 발생했습니다: {str(e)}")
