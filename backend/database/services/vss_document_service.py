"""문서(Word/PDF) 생성 관련 서비스 (클래스 기반)
"""
import logging
import io
import tempfile
import time
from typing import List, Optional, Tuple
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image
import base64

from app_config.settings import REPORTS_DIR
from database.services.vss_thumbnail_service import ThumbnailService

logger = logging.getLogger(__name__)


class DocumentService:
    @staticmethod
    def add_horizontal_line(paragraph):
        p = paragraph._element
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        pPr.insert_element_before(pBdr, 'w:shd')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)
        return paragraph

    @staticmethod
    def create_word_document(user_id: str, title: str, author: str, description: str, query: str, clips: List[dict]) -> Tuple[Document, str, int, List[Optional[str]]]:
        """
        클립 목록으로 Word Document 객체를 생성하고, 데이터베이스 저장용 텍스트와 단어 수를 반환합니다.
        반환값: (Document 객체, report_content(str), word_count(int), list_of_base64_thumbnails)
        """
        doc = Document()
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        title_heading = doc.add_heading(title, level=1)
        title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_heading.runs[0]
        title_run.font.size = Pt(24)
        title_run.bold = True
        title_run.font.color.rgb = RGBColor(0, 51, 102)

        doc.add_paragraph(); doc.add_paragraph()
        separator_para = doc.add_paragraph(); DocumentService.add_horizontal_line(separator_para)
        doc.add_paragraph(); doc.add_paragraph()

        current_date = time.strftime('%Y년 %m월 %d일')
        info_table = doc.add_table(rows=2, cols=2)
        info_table.style = 'Light Grid Accent 1'
        info_table.rows[0].cells[0].text = '작성자'
        info_table.rows[0].cells[1].text = author
        info_table.rows[1].cells[0].text = '작성 일자'
        info_table.rows[1].cells[1].text = current_date

        doc.add_paragraph(); doc.add_paragraph()
        DocumentService.add_horizontal_line(doc.add_paragraph())
        doc.add_paragraph(); doc.add_paragraph()

        query_section_para = doc.add_paragraph()
        query_label = query_section_para.add_run('Q. 검색어')
        query_label.bold = True
        query_label.font.size = Pt(16)
        query_label.font.color.rgb = RGBColor(0, 102, 204)
        doc.add_paragraph()
        query_content_para = doc.add_paragraph()
        query_content_para.paragraph_format.left_indent = Inches(0.5)
        query_content_run = query_content_para.add_run(query if query else '검색어가 제공되지 않았습니다.')
        query_content_run.font.size = Pt(13)
        doc.add_paragraph(); doc.add_paragraph()
        DocumentService.add_horizontal_line(doc.add_paragraph())

        # A. 검색 결과
        answer_para = doc.add_paragraph()
        answer_label = answer_para.add_run('A. 검색 결과')
        answer_label.bold = True
        answer_label.font.size = Pt(16)
        answer_label.font.color.rgb = RGBColor(0, 102, 204)
        doc.add_paragraph()

        report_content = f"# {title}\n\n**작성자:** {author}\n\n**작성 일자:** {current_date}\n\n" + ("="*50) + "\n\n## Q. 검색어\n\n" + (query if query else '검색어가 제공되지 않았습니다.') + "\n\n" + ("="*50) + "\n\n## A. 검색 결과\n\n"

        thumb_base64_list = []
        for idx, clip in enumerate(clips, 1):
            if idx > 1:
                clip_separator = doc.add_paragraph(); DocumentService.add_horizontal_line(clip_separator); doc.add_paragraph()

            clip_para = doc.add_paragraph()
            clip_para.paragraph_format.left_indent = Inches(0.3)
            clip_number = clip_para.add_run(f"[{idx}] ")
            clip_number.bold = True
            clip_number.font.size = Pt(13)

            sentence = clip.get('sentence') or clip.get('title')
            if sentence:
                clip_para.add_run(sentence)

            if clip.get('start_time') is not None and clip.get('end_time') is not None:
                time_info = f" (시간: {clip['start_time']:.2f}초 - {clip['end_time']:.2f}초)"
                clip_para.add_run(time_info)

            # 썸네일
            image_data = None
            if clip.get('sourceVideo') and clip.get('start_time') is not None:
                try:
                    original_video_path = ThumbnailService.get_original_video_path(clip.get('sourceVideo'), user_id)
                    if original_video_path:
                        image_data = ThumbnailService.get_video_thumbnail(original_video_path, clip.get('start_time'))
                except Exception:
                    pass

            if not image_data and clip.get('url'):
                try:
                    if clip.get('url').lower().endswith(('.jpg', '.jpeg', '.png', '.gif', '.webp')):
                        image_data = ThumbnailService.download_image(clip.get('url'))
                    else:
                        thumbnail_time = clip.get('start_time') if clip.get('start_time') is not None else 0.0
                        image_data = ThumbnailService.get_video_thumbnail(clip.get('url'), thumbnail_time)
                except Exception:
                    pass

            if image_data:
                try:
                    image = Image.open(io.BytesIO(image_data))
                    max_width = 5.0
                    width, height = image.size
                    if width > max_width * 72:
                        ratio = (max_width * 72) / width
                        new_width = int(width * ratio)
                        new_height = int(height * ratio)
                        image = image.resize((new_width, new_height), Image.Resampling.LANCZOS)
                    img_byte_arr = io.BytesIO()
                    image.save(img_byte_arr, format='PNG')
                    img_byte_arr.seek(0)
                    run = doc.add_paragraph().add_run()
                    try:
                        run.add_picture(img_byte_arr, width=Inches(max_width))
                    except Exception:
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_img:
                            tmp_img.write(img_byte_arr.getvalue())
                            run.add_picture(tmp_img.name, width=Inches(max_width))
                    img_byte_arr.seek(0)
                    thumb_base64_list.append(base64.b64encode(img_byte_arr.read()).decode('utf-8'))
                except Exception:
                    thumb_base64_list.append(None)
            else:
                thumb_base64_list.append(None)

            doc.add_paragraph()

            # report_content
            if sentence:
                report_content += f"{idx}. {sentence}"
                if clip.get('start_time') is not None and clip.get('end_time') is not None:
                    report_content += f" ({clip['start_time']:.2f}초 - {clip['end_time']:.2f}초)"
                report_content += "\n\n"

        word_count = len(report_content.split())
        return doc, report_content, word_count, thumb_base64_list

    @staticmethod
    def convert_docx_to_pdf(docx_path: str, pdf_path: Optional[str] = None) -> Optional[str]:
        # 간단 래퍼 — 기존 환경에서 docx2pdf가 설정되어 있지 않을 수 있음.
        try:
            from docx2pdf import convert
        except Exception:
            logger.warning('docx2pdf 미설치: PDF 변환 불가')
            return None
        try:
            if pdf_path is None:
                pdf_path = str(Path(docx_path).with_suffix('.pdf'))
            convert(docx_path, pdf_path)
            return pdf_path if Path(pdf_path).exists() else None
        except Exception as e:
            logger.warning(f"PDF 변환 실패: {e}")
            return None
