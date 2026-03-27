"""보고서 관련 라우터"""
import json
import logging
import os
import time
import glob
import re
from datetime import datetime
from typing import Optional, List
from pathlib import Path
from fastapi import APIRouter, HTTPException, Query
from fastapi.responses import FileResponse
from pydantic import BaseModel
import io
import tempfile
import base64
from docx import Document
from docx.shared import Inches, Pt
from PIL import Image

from routers.auth import verify_user_exists
from dependencies import verify_user_dependency
from exceptions import NotFoundException, ValidationException, DatabaseException

# 서비스 레이어 분리: 썸네일/문서/파일/리포트 오케스트레이터
# ReportService 클래스 기반 서비스 사용
from database.services.vss_reports_service import ReportService
from database.services.vss_file_service import FileService
from database.services.vss_thumbnail_service import ThumbnailService
from database.services.vss_document_service import DocumentService
from app_config.settings import REPORTS_DIR
from fastapi import Depends

logger = logging.getLogger(__name__)

router = APIRouter()

# ==================== 요청 모델 ====================
class CreateReportRequest(BaseModel):
    user_id: str
    title: str
    description: Optional[str] = None
    content: str
    word_count: int = 0
    video_ids: Optional[List[int]] = None
    video_titles: Optional[List[str]] = None

class CreateReportResponse(BaseModel):
    success: bool
    report_id: Optional[int] = None
    message: str

class ClipData(BaseModel):
    id: Optional[str] = None
    title: str
    url: str
    sentence: Optional[str] = None
    start_time: Optional[float] = None
    end_time: Optional[float] = None
    sourceVideo: Optional[str] = None

class CreateWordReportRequest(BaseModel):
    user_id: str
    title: str
    author: Optional[str] = None
    description: Optional[str] = None
    query: Optional[str] = None
    clips: List[ClipData]

class AddClipsToReportRequest(BaseModel):
    user_id: str
    clips: List[ClipData]

class UpdateReportRequest(BaseModel):
    user_id: str
    title: Optional[str] = None
    description: Optional[str] = None
    content: Optional[str] = None

@router.post("", response_model=CreateReportResponse)
async def create_report(request: CreateReportRequest):
    logger.info("[VSS] create_report")
    """보고서 생성"""
    try:
        user_id = request.user_id
        title = request.title
        description = request.description or ""
        content = request.content
        word_count = request.word_count or 0
        video_ids = request.video_ids or []
        video_titles = request.video_titles or []
        
        if not user_id:
            raise HTTPException(status_code=400, detail="사용자 ID가 필요합니다.")
        if not title:
            raise HTTPException(status_code=400, detail="보고서 제목이 필요합니다.")
        if not content:
            raise HTTPException(status_code=400, detail="보고서 내용이 필요합니다.")
        
        verify_user_exists(user_id)
        
        # 보고서 저장
        video_ids_json = json.dumps(video_ids) if video_ids else None
        video_titles_json = json.dumps(video_titles) if video_titles else None

        logger.info(f"video_ids_json={video_ids_json}, video_titles_json={video_titles_json}")

        report_id = ReportService.create_report(user_id, title, description, content, word_count, video_ids_json, video_titles_json)
        
        logger.info(f"보고서 생성 완료: USER_ID={user_id}, REPORT_ID={report_id}, TITLE={title}")
        
        return {
            "success": True,
            "report_id": report_id,
            "message": "보고서가 성공적으로 생성되었습니다."
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"보고서 생성 실패: {e}")
        raise HTTPException(status_code=500, detail=f"보고서 생성 중 오류가 발생했습니다: {str(e)}")

@router.post("/create-word")
async def create_word_report(request: CreateWordReportRequest):
    logger.info("[VSS] create_word_report")
    try:
        if not request.user_id:
            raise HTTPException(status_code=400, detail="사용자 ID가 필요합니다.")
        if not request.title:
            raise HTTPException(status_code=400, detail="보고서 제목이 필요합니다.")
        if not request.clips or len(request.clips) == 0:
            raise HTTPException(status_code=400, detail="클립 데이터가 필요합니다.")

        verify_user_exists(request.user_id)

        # 서비스 레이어에 위임 (문서 생성 -> DB 저장 -> 파일 저장)
        res = ReportService.create_word_report(
            request.user_id,
            request.title,
            request.author or request.user_id,
            request.description or "",
            request.query or "",
            [c.dict() for c in request.clips]
        )

        return {
            "success": True,
            "report_id": res.get('report_id'),
            "file_url": res.get('file_url'),
            "message": "보고서가 성공적으로 생성되었습니다."
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"워드 보고서 생성 실패: {e}")
        raise HTTPException(status_code=500, detail=f"보고서 생성 중 오류가 발생했습니다: {str(e)}")

@router.get("")
async def get_reports(
    user_id: str = Query(..., description="사용자 ID"),
    page: int = Query(1, ge=1, description="페이지 번호"),
    page_size: int = Query(10, ge=1, le=100, description="페이지당 항목 수")
):
    logger.info("[VSS] get_reports")
    """보고서 목록 조회 (페이지네이션 지원)"""
    try:
        if not user_id:
            raise HTTPException(status_code=400, detail="사용자 ID가 필요합니다.")
        
        verify_user_exists(user_id)
        
        rows, total = ReportService.get_reports(user_id, page, page_size)
        pages = max(1, (total + page_size - 1) // page_size) if total > 0 else 0
        
        # context manager 밖에서 데이터 처리
        reports = []
        for row in rows:
            report_id, title, description, content, word_count, video_ids_json, video_titles_json, created_at, updated_at = row
            
            # Word 파일 경로 생성 (보고서 ID 기반)
            # 실제 파일 찾기
            pattern = str(REPORTS_DIR / f"report_{report_id}_*.docx")
            matching_files = glob.glob(pattern)
            if matching_files:
                filename = os.path.basename(matching_files[0])
                file_url = f"/reports-files/{filename}"
            else:
                file_url = None
            
            # JSON 문자열 파싱
            video_ids = json.loads(video_ids_json) if video_ids_json else []
            video_titles = json.loads(video_titles_json) if video_titles_json else []
            
            reports.append({
                "id": report_id,
                "report_id": report_id,
                "title": title,
                "description": description or "",
                "content": content or "",
                "word_count": word_count or 0,
                "video_ids": video_ids,
                "video_titles": video_titles,
                "file_url": file_url,
                "created_at": created_at.isoformat() if created_at else None,
                "createdAt": created_at.isoformat() if created_at else None,
                "updated_at": updated_at.isoformat() if updated_at else None
            })
        
        logger.info(f"보고서 목록 조회 완료: USER_ID={user_id}, 총 {total}개, 페이지 {page}/{pages}")
        
        return {
            "success": True,
            "reports": reports,
            "total": total,
            "page": page,
            "page_size": page_size,
            "pages": pages
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"보고서 목록 조회 실패: {e}")
        raise HTTPException(status_code=500, detail=f"보고서 목록 조회 중 오류가 발생했습니다: {str(e)}")


@router.get("/{report_id}")
async def get_report(

    report_id: int,
    user_id: str = Query(..., description="사용자 ID")
):
    logger.info("[VSS] get_report")
    """보고서 상세 조회"""
    try:
        if not user_id:
            raise HTTPException(status_code=400, detail="사용자 ID가 필요합니다.")
        
        verify_user_exists(user_id)
        
        row = ReportService.get_report(report_id, user_id)
        
        if not row:
            raise HTTPException(status_code=404, detail="보고서를 찾을 수 없습니다.")
        
        report_id_db, title, description, content, word_count, video_ids_json, video_titles_json, created_at, updated_at = row
        
        # JSON 문자열 파싱
        video_ids = json.loads(video_ids_json) if video_ids_json else []
        video_titles = json.loads(video_titles_json) if video_titles_json else []
        
        # Word 파일 경로 찾기
        pattern = str(REPORTS_DIR / f"report_{report_id_db}_*.docx")
        matching_files = glob.glob(pattern)
        if matching_files:
            filename = os.path.basename(matching_files[0])
            file_url = f"/reports-files/{filename}"
        else:
            file_url = None
        
        return {
            "success": True,
            "report": {
                "id": report_id_db,
                "report_id": report_id_db,
                "title": title,
                "description": description or "",
                "content": content or "",
                "word_count": word_count or 0,
                "video_ids": video_ids,
                "video_titles": video_titles,
                "file_url": file_url,
                "created_at": created_at.isoformat() if created_at else None,
                "createdAt": created_at.isoformat() if created_at else None,
                "updated_at": updated_at.isoformat() if updated_at else None
            }
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"보고서 상세 조회 실패: {e}")
        raise HTTPException(status_code=500, detail=f"보고서 상세 조회 중 오류가 발생했습니다: {str(e)}")

@router.delete("/{report_id}")
async def delete_report(
    report_id: int,
    user_id: str = Query(..., description="사용자 ID")
):
    logger.info("[VSS] delete_report")
    """보고서 삭제"""
    try:
        if not user_id:
            raise HTTPException(status_code=400, detail="사용자 ID가 필요합니다.")
        
        verify_user_exists(user_id)
        
        existing = ReportService.get_report(report_id, user_id)
        if not existing:
            raise HTTPException(status_code=404, detail="보고서를 찾을 수 없거나 권한이 없습니다.")
        
        # Word 파일 삭제 시도
        pattern = str(REPORTS_DIR / f"report_{report_id}_*.docx")
        matching_files = glob.glob(pattern)
        for file_path in matching_files:
            try:
                if os.path.exists(file_path):
                    os.unlink(file_path)
                    logger.info(f"Word 파일 삭제 완료: {file_path}")
            except Exception as e:
                logger.warning(f"Word 파일 삭제 실패 ({file_path}): {e}")
        
        # 보고서 삭제
        deleted_count = ReportService.delete_report(report_id, user_id)
        
        if deleted_count == 0:
            # 이미 삭제되었을 수 있음
            logger.warning(f"보고서 삭제 실패: USER_ID={user_id}, REPORT_ID={report_id} (이미 삭제되었을 수 있음)")
            return {
                "success": True,
                "message": "보고서가 이미 삭제되었거나 존재하지 않습니다."
            }
        
        logger.info(f"보고서 삭제 완료: USER_ID={user_id}, REPORT_ID={report_id}")
        
        return {
            "success": True,
            "message": "보고서가 성공적으로 삭제되었습니다."
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"보고서 삭제 실패: {e}")
        raise HTTPException(status_code=500, detail=f"보고서 삭제 중 오류가 발생했습니다: {str(e)}")

@router.put("/{report_id}/add-clips")
async def add_clips_to_report(
    report_id: int,
    request: AddClipsToReportRequest
):
    logger.info("add_clips_to_report")
    """기존 보고서에 클립 추가"""
    try:
        user_id = request.user_id
        clips = request.clips
        
        if not user_id:
            raise HTTPException(status_code=400, detail="사용자 ID가 필요합니다.")
        if not clips or len(clips) == 0:
            raise HTTPException(status_code=400, detail="클립 데이터가 필요합니다.")
        
        verify_user_exists(user_id)
        
        row = ReportService.get_report(report_id, user_id)
        if not row:
            raise HTTPException(status_code=404, detail="보고서를 찾을 수 없거나 권한이 없습니다.")
        report_id_db, title, description, existing_content, word_count, video_ids_json, video_titles_json, _, _ = row
        
        # 기존 클립 개수 계산 (콘텐츠에서 ## 숫자. 패턴으로 찾기)
        existing_clip_count = len(re.findall(r'## \d+\.', existing_content))
        
        # 기존 클립 ID와 URL 추출 (중복 체크용)
        existing_video_ids = json.loads(video_ids_json) if video_ids_json else []
        existing_clip_urls = set()
        
        # 기존 콘텐츠에서 클립 URL 추출 (마크다운 형식에서)
        url_pattern = r'- 비디오 URL: (https?://[^\s]+)'
        existing_urls = re.findall(url_pattern, existing_content)
        existing_clip_urls.update(existing_urls)
        
        # 기존 콘텐츠에서 클립 제목도 추출 (제목으로도 중복 체크)
        existing_clip_titles = set()
        title_pattern = r'## \d+\. ([^\n]+)'
        existing_titles = re.findall(title_pattern, existing_content)
        existing_clip_titles.update([t.strip() for t in existing_titles])
        
        # 중복 클립 필터링 (URL 또는 ID 기준)
        new_clips = []
        duplicate_clips = []
        
        for clip in clips:
            is_duplicate = False
            
            # URL로 중복 체크
            if clip.url and clip.url in existing_clip_urls:
                is_duplicate = True
            
            # ID로 중복 체크
            if clip.id:
                try:
                    clip_id = int(clip.id)
                    if clip_id in existing_video_ids:
                        is_duplicate = True
                except (ValueError, TypeError):
                    pass
            
            # 제목으로도 중복 체크 (URL과 ID가 없는 경우)
            if not is_duplicate and clip.title and clip.title.strip() in existing_clip_titles:
                is_duplicate = True
            
            if is_duplicate:
                duplicate_clips.append(clip)
            else:
                new_clips.append(clip)
        
        # 중복 클립이 모두인 경우
        if len(new_clips) == 0:
            return {
                "success": False,
                "message": "추가할 수 있는 새로운 클립이 없습니다. 모든 클립이 이미 보고서에 포함되어 있습니다.",
                "duplicate_count": len(duplicate_clips)
            }
        
        # 일부만 중복인 경우 경고 메시지 포함
        duplicate_message = ""
        if len(duplicate_clips) > 0:
            duplicate_message = f" ({len(duplicate_clips)}개의 중복 클립은 제외되었습니다)"
        
        # 새 클립만 사용
        clips = new_clips
        
        # 기존 Word 파일 찾기
        pattern = str(REPORTS_DIR / f"report_{report_id}_*.docx")
        matching_files = glob.glob(pattern)
        if not matching_files:
            raise HTTPException(status_code=404, detail="Word 파일을 찾을 수 없습니다.")
        
        word_file_path = matching_files[0]
        
        # 기존 Word 문서 열기
        doc = Document(str(word_file_path))
        
        # 기존 클립 수 확인 (제목 개수로 추정)
        existing_clip_count = len([p for p in doc.paragraphs if p.style.name == 'Heading 1' and p.text.strip() and not p.text.strip().startswith('#')])
        
        # 새로운 클립 추가
        clip_images = []
        new_clip_start_idx = existing_clip_count + 1
        
        for idx, clip in enumerate(clips, new_clip_start_idx):
            # 구분선 추가 (이전 클립과 구분)
            separator_para = doc.add_paragraph()
            DocumentService.add_horizontal_line(separator_para)
            
            # 클립 번호 및 제목
            doc.add_heading(f"{idx}. {clip.title}", level=1)
            
            # 시간 정보
            if clip.start_time is not None and clip.end_time is not None:
                time_para = doc.add_paragraph(f"시간: {format_time(clip.start_time)} - {format_time(clip.end_time)}")
                time_para.style = 'List Bullet'
            
            # 소스 비디오 정보
            if clip.sourceVideo:
                source_para = doc.add_paragraph(f"소스: {clip.sourceVideo}")
                source_para.style = 'List Bullet'
            
            # 썸네일 이미지 추가 시도
            # 원본 동영상에서 시작 타임스탬프 부분을 썸네일로 사용
            image_data = None
            if clip.sourceVideo and clip.start_time is not None:
                try:
                    logger.info(f"원본 동영상 경로 찾기 시도: sourceVideo={clip.sourceVideo}, start_time={clip.start_time}, user_id={user_id}")
                    # 원본 동영상 경로 찾기
                    original_video_path = ThumbnailService.get_original_video_path(clip.sourceVideo, user_id)
                    if original_video_path:
                        logger.info(f"✅ 원본 동영상 경로 찾음: {original_video_path}")
                        logger.info(f"원본 동영상에서 썸네일 추출: {original_video_path}, 시간: {clip.start_time}")
                        image_data = ThumbnailService.get_video_thumbnail(original_video_path, clip.start_time)
                        if image_data:
                            logger.info(f"✅ 원본 동영상에서 썸네일 추출 성공: {len(image_data)} bytes")
                        else:
                            logger.warning("⚠️ 원본 동영상에서 썸네일 추출 실패 (None 반환)")
                    else:
                        logger.warning(f"⚠️ 원본 동영상을 찾을 수 없음: {clip.sourceVideo}")
                except Exception as e:
                    logger.warning(f"원본 동영상 썸네일 추출 실패 ({clip.sourceVideo}): {e}")
                    import traceback
                    logger.warning(f"원본 동영상 썸네일 추출 실패 상세: {traceback.format_exc()}")
            
            # 원본 동영상에서 썸네일을 가져오지 못한 경우 기존 로직 사용 (fallback)
            if not image_data and clip.url:
                try:
                    if clip.url.lower().endswith(('.jpg', '.jpeg', '.png', '.gif', '.webp')):
                        image_data = ThumbnailService.download_image(clip.url)
                    else:
                        thumbnail_time = clip.start_time if clip.start_time is not None else 0.0
                        image_data = ThumbnailService.get_video_thumbnail(clip.url, thumbnail_time)
                except Exception as e:
                    logger.warning(f"썸네일 처리 실패 ({clip.url}): {e}")
            
            if image_data:
                try:
                    logger.info(f"썸네일 이미지 데이터 수신: {len(image_data)} bytes (클립 추가)")
                    image = Image.open(io.BytesIO(image_data))
                    max_width = 4.0
                    width, height = image.size
                    logger.info(f"이미지 크기: {width}x{height} (클립 추가)")
                    if width > max_width * 72:
                        ratio = (max_width * 72) / width
                        new_width = int(width * ratio)
                        new_height = int(height * ratio)
                        image = image.resize((new_width, new_height), Image.Resampling.LANCZOS)
                        logger.info(f"이미지 리사이즈: {new_width}x{new_height} (클립 추가)")
                    
                    img_byte_arr = io.BytesIO()
                    image.save(img_byte_arr, format='PNG')
                    img_byte_arr.seek(0)
                    try:
                        doc.add_picture(img_byte_arr, width=Inches(max_width))
                        logger.info("✅ 썸네일 이미지 추가 성공 (클립 추가)")
                    except Exception as pic_err:
                        logger.warning(f"doc.add_picture(BytesIO) 실패, 임시 파일로 재시도: {pic_err}")
                        try:
                            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_img:
                                tmp_img.write(img_byte_arr.getvalue())
                                tmp_path = tmp_img.name
                            doc.add_picture(tmp_path, width=Inches(max_width))
                            logger.info("✅ 썸네일 이미지 추가 성공 (임시파일, 클립 추가)")
                        except Exception as tmp_err:
                            logger.warning(f"임시 파일로도 이미지 추가 실패 (클립 추가): {tmp_err}")
                        finally:
                            try:
                                if 'tmp_path' in locals() and os.path.exists(tmp_path):
                                    os.unlink(tmp_path)
                            except Exception:
                                pass

                    img_byte_arr.seek(0)
                    clip_images.append(base64.b64encode(img_byte_arr.read()).decode('utf-8'))
                except Exception as e:
                    logger.warning(f"이미지 추가 실패: {e}")
                    import traceback
                    logger.warning(f"이미지 추가 실패 상세: {traceback.format_exc()}")
                    clip_images.append(None)
            else:
                logger.warning("⚠️ 썸네일 이미지 데이터 없음 (클립 추가)")
                clip_images.append(None)
            
            # 장면 설명 추가
            if clip.sentence:
                sentence_para = doc.add_paragraph("장면 설명:")
                sentence_para.style = 'Heading 3'
                desc_para = doc.add_paragraph(clip.sentence)
                desc_para.style = 'Normal'
        
        # 기존 내용에 새 클립 내용 추가
        new_content = existing_content
        for idx, clip in enumerate(clips, new_clip_start_idx):
            new_content += f"\n## {idx}. {clip.title}\n\n"
            if clip.start_time is not None and clip.end_time is not None:
                new_content += f"- 시간: {format_time(clip.start_time)} - {format_time(clip.end_time)}\n"
            if clip.sourceVideo:
                new_content += f"- 소스: {clip.sourceVideo}\n"
            
            clip_image_idx = idx - new_clip_start_idx
            if clip_image_idx < len(clip_images) and clip_images[clip_image_idx]:
                new_content += f"\n![썸네일 {idx}](data:image/png;base64,{clip_images[clip_image_idx]})\n\n"
            elif clip.url:
                new_content += f"- 비디오 URL: {clip.url}\n"
            
            if clip.sentence:
                new_content += f"\n### 장면 설명:\n{clip.sentence}\n\n"
            new_content += "=" * 50 + "\n\n"
        
        # 단어 수 재계산
        word_count = len(new_content.split())
        
        # 기존 VIDEO_IDS와 VIDEO_TITLES에 새 클립 추가
        existing_video_ids = json.loads(video_ids_json) if video_ids_json else []
        existing_video_titles = json.loads(video_titles_json) if video_titles_json else []
        
        for clip in clips:
            if clip.id:
                try:
                    clip_id = int(clip.id)
                    if clip_id not in existing_video_ids:
                        existing_video_ids.append(clip_id)
                except (ValueError, TypeError):
                    pass
            if clip.title and clip.title not in existing_video_titles:
                existing_video_titles.append(clip.title)
        
        video_ids_json = json.dumps(existing_video_ids) if existing_video_ids else None
        video_titles_json = json.dumps(existing_video_titles) if existing_video_titles else None
        
        # description 업데이트 (클립 개수 반영)
        # 기존 클립 수와 새로 추가된 클립 수를 합산
        total_clip_count = existing_clip_count + len(clips)
        
        # description이 "n개의 클립 검색 결과" 형식인지 확인
        description_pattern_ko = r'(\d+)개의 클립 검색 결과'
        description_pattern_en = r'Search results for (\d+) clips?'
        
        updated_description = description
        if description:
            # 한국어 패턴 매칭
            match_ko = re.search(description_pattern_ko, description)
            if match_ko:
                updated_description = re.sub(description_pattern_ko, f'{total_clip_count}개의 클립 검색 결과', description)
            else:
                # 영어 패턴 매칭
                match_en = re.search(description_pattern_en, description)
                if match_en:
                    updated_description = re.sub(description_pattern_en, f'Search results for {total_clip_count} clips', description)
                else:
                    # 패턴이 없으면 description에 클립 개수 정보 추가
                    # 한국어/영어 판단
                    is_korean = any(ord(c) >= 0xAC00 and ord(c) <= 0xD7A3 for c in description) if description else True
                    if is_korean:
                        updated_description = f"{total_clip_count}개의 클립 검색 결과"
                    else:
                        updated_description = f"Search results for {total_clip_count} clips"
        else:
            # description이 없으면 새로 생성 (기본값은 한국어)
            updated_description = f"{total_clip_count}개의 클립 검색 결과"
        
        # Word 파일의 description도 업데이트
        # Word 문서에서 description 부분 찾기 및 업데이트
        if doc.paragraphs:
            # 제목 다음 단락들을 확인 (description은 보통 제목 바로 다음)
            for i, para in enumerate(doc.paragraphs):
                para_text = para.text.strip()
                # description 패턴이 있는 단락 찾기
                if re.search(description_pattern_ko, para_text) or re.search(description_pattern_en, para_text):
                    # 단락 텍스트 업데이트
                    new_text = re.sub(description_pattern_ko, f'{total_clip_count}개의 클립 검색 결과', para_text)
                    new_text = re.sub(description_pattern_en, f'Search results for {total_clip_count} clips', new_text)
                    if new_text != para_text:
                        para.clear()
                        para.add_run(new_text)
                    break
                # 제목 다음 3개 단락 내에서만 확인 (성능 최적화)
                if i > 5:  # 제목, 작성자, 구분선 등을 고려하여 적절한 범위
                    break
        
        # Word 파일 저장
        doc.save(str(word_file_path))
        
        # 데이터베이스 업데이트 (description 포함)
        # Update via service (repository/ORM)
        ReportService.update_report_content(report_id, user_id, new_content, updated_description, word_count, video_ids_json, video_titles_json)
        
        logger.info(f"보고서에 클립 추가 완료: USER_ID={user_id}, REPORT_ID={report_id}, 추가된 클립 수={len(clips)}")
        
        return {
            "success": True,
            "report_id": report_id,
            "file_url": f"/reports-files/{os.path.basename(word_file_path)}",
            "message": f"{len(clips)}개의 클립이 보고서에 추가되었습니다.{duplicate_message}",
            "added_count": len(clips),
            "duplicate_count": len(duplicate_clips) if 'duplicate_clips' in locals() else 0
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"보고서에 클립 추가 실패: {e}")
        raise HTTPException(status_code=500, detail=f"보고서에 클립 추가 중 오류가 발생했습니다: {str(e)}")

@router.put("/{report_id}")
async def update_report(
    report_id: int,
    request: UpdateReportRequest
):
    logger.info("[VSS] update_report")
    """보고서 내용 수정"""
    try:
        user_id = request.user_id
        
        if not user_id:
            raise HTTPException(status_code=400, detail="사용자 ID가 필요합니다.")
        
        verify_user_exists(user_id)
        
        row = get_report(report_id, user_id)
        if not row:
            raise HTTPException(status_code=404, detail="보고서를 찾을 수 없거나 권한이 없습니다.")
        report_id_db, existing_title, existing_description, existing_content, existing_word_count = row
        
        # 업데이트할 필드 준비
        title = request.title if request.title is not None else existing_title
        description = request.description if request.description is not None else existing_description
        content = request.content if request.content is not None else existing_content
        
        # 단어 수 재계산 (내용이 변경된 경우)
        word_count = existing_word_count
        if request.content is not None:
            word_count = len(content.split())
        
        # 기존 Word 파일 경로 확인 (파일 시스템에서 찾기)
        pattern = str(REPORTS_DIR / f"report_{report_id}_*.docx")
        matching_files = glob.glob(pattern)
        existing_file_url = None
        if matching_files:
            filename = os.path.basename(matching_files[0])
            existing_file_url = f"/reports-files/{filename}"
        
        # Word 파일 재생성 (내용이 변경된 경우)
        new_file_url = existing_file_url
        if request.content is not None and content != existing_content:
            try:
                # 마크다운 내용을 Word 문서로 변환
                doc = Document()
            
                lines = content.split('\n')
                current_paragraph = None
                in_list = False
                
                for line in lines:
                    line = line.strip()
                    
                    # 빈 줄 처리
                    if not line:
                        if current_paragraph and not in_list:
                            current_paragraph = None
                        in_list = False
                        continue
                    
                    # 제목 처리 (# ## ###)
                    if line.startswith('#'):
                        in_list = False
                        level = len(line) - len(line.lstrip('#'))
                        text = line.lstrip('#').strip()
                        if text:
                            doc.add_heading(text, level=min(level, 3))
                            current_paragraph = None
                    # 이미지 처리 (![alt](data:image/...))
                    elif '![' in line and 'data:image' in line:
                        in_list = False
                        # base64 이미지 추출 (여러 줄에 걸쳐 있을 수 있음)
                        match = re.search(r'data:image/([^;]+);base64,([A-Za-z0-9+/=]+)', line)
                        if match:
                            _, img_data = match.groups()
                            try:
                                img_bytes = base64.b64decode(img_data)
                                img_io = io.BytesIO(img_bytes)
                                doc.add_picture(img_io, width=Inches(4))
                                current_paragraph = None
                            except Exception as e:
                                logger.warning(f"이미지 추가 실패: {e}")
                    # 리스트 항목 처리 (- * 숫자.)
                    elif re.match(r'^[-*]\s+', line) or re.match(r'^\d+\.\s+', line):
                        in_list = True
                        text = re.sub(r'^[-*]\s+', '', line)
                        text = re.sub(r'^\d+\.\s+', '', text)
                        if text:
                            para = doc.add_paragraph(text, style='List Bullet')
                            current_paragraph = para
                    # 구분선 처리 (=== 또는 ---)
                    elif re.match(r'^={3,}$', line) or re.match(r'^-{3,}$', line):
                        in_list = False
                        separator_para = doc.add_paragraph()
                        DocumentService.add_horizontal_line(separator_para)
                        current_paragraph = None
                    # 볼드/이탤릭 처리 (**text** 또는 *text*)
                    elif '**' in line or '*' in line:
                        in_list = False
                        # 간단한 볼드 처리
                        text = line.replace('**', '').replace('*', '')
                        if text:
                            para = doc.add_paragraph()
                            para.add_run(text)
                            current_paragraph = para
                    # 일반 텍스트
                    else:
                        if in_list:
                            in_list = False
                        if current_paragraph and not in_list:
                            current_paragraph.add_run(' ' + line)
                        else:
                            current_paragraph = doc.add_paragraph(line)
                
                # Word 파일 저장
                REPORTS_DIR.mkdir(exist_ok=True, parents=True)
                timestamp = int(time.time() * 1000)
                final_filename = f"report_{report_id}_{timestamp}.docx"
                final_file_path = REPORTS_DIR / final_filename
                doc.save(str(final_file_path))
                
                # 기존 파일 삭제 (있는 경우)
                if existing_file_url:
                    try:
                        old_filename = existing_file_url.replace("/reports-files/", "")
                        old_file_path = REPORTS_DIR / old_filename
                        if old_file_path.exists() and old_file_path != final_file_path:
                            old_file_path.unlink()
                            logger.info(f"기존 Word 파일 삭제: {old_filename}")
                    except Exception as e:
                        logger.warning(f"기존 Word 파일 삭제 실패: {e}")
                
                new_file_url = f"/reports-files/{final_filename}"
                logger.info(f"Word 파일 재생성 완료: {final_filename}")
            except Exception as e:
                logger.error(f"Word 파일 재생성 실패: {e}")
                # Word 파일 재생성 실패해도 DB 업데이트는 계속 진행
        
        # 데이터베이스 업데이트
        ReportService.update_report(report_id, user_id, title, description, content, word_count)
        
        logger.info(f"보고서 수정 완료: USER_ID={user_id}, REPORT_ID={report_id}")
        
        return {
            "success": True,
            "report_id": report_id,
            "file_url": new_file_url,
            "message": "보고서가 성공적으로 수정되었습니다."
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"보고서 수정 실패: {e}")
        raise HTTPException(status_code=500, detail=f"보고서 수정 중 오류가 발생했습니다: {str(e)}")

def format_time(seconds: Optional[float]) -> str:
    logger.info(f"[VSS] format_time 호출: seconds={seconds}")
    """초를 MM:SS 형식으로 변환"""
    if seconds is None:
        return ""
    minutes = int(seconds // 60)
    secs = int(seconds % 60)
    return f"{minutes:02d}:{secs:02d}"

# 썸네일/문서/파일 서비스는 `backend/database/services`로 이동했습니다.

# PDF 변환은 document service로 이동했습니다 (vss_document_service.convert_docx_to_pdf)

@router.get("/{report_id}/pdf")
async def get_report_pdf(
    report_id: int,
    user_id: str = Query(..., description="사용자 ID")
):
    logger.info(f"[VSS] get_report_pdf 호출: report_id={report_id}, user_id={user_id}")
    """보고서를 PDF로 변환하여 반환"""
    try:
        if not user_id:
            raise HTTPException(status_code=400, detail="사용자 ID가 필요합니다.")
        
        # PDF 변환은 document service에서 처리합니다.
        verify_user_exists(user_id)
        
        row = get_report(report_id, user_id)
        if not row:
            raise HTTPException(status_code=404, detail="보고서를 찾을 수 없거나 권한이 없습니다.")
        report_id_db, title, _, _, _, _, file_url, _, _ = row
        
        # Word 파일 경로 찾기
        pattern = str(REPORTS_DIR / f"report_{report_id_db}_*.docx")
        matching_files = glob.glob(pattern)
        if not matching_files:
            raise HTTPException(status_code=404, detail="Word 파일을 찾을 수 없습니다.")
        
        docx_file_path = matching_files[0]
        
        # PDF 파일 경로 생성
        pdf_file_path = str(Path(docx_file_path).with_suffix('.pdf'))
        
        # PDF 파일이 없거나 DOCX 파일이 더 최신인 경우 변환
        docx_mtime = os.path.getmtime(docx_file_path)
        pdf_exists = os.path.exists(pdf_file_path)
        pdf_mtime = os.path.getmtime(pdf_file_path) if pdf_exists else 0
        
        if not pdf_exists or docx_mtime > pdf_mtime:
            # PDF 변환 실행
            converted_pdf_path = DocumentService.convert_docx_to_pdf(docx_file_path, pdf_file_path)
            if not converted_pdf_path:
                raise HTTPException(status_code=500, detail="PDF 변환에 실패했습니다.")
        
        # PDF 파일 반환
        if not os.path.exists(pdf_file_path):
            raise HTTPException(status_code=404, detail="PDF 파일을 찾을 수 없습니다.")
        
        return FileResponse(
            pdf_file_path,
            media_type="application/pdf",
            filename=f"{title}_{report_id_db}.pdf"
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"PDF 변환 실패: {e}")
        raise HTTPException(status_code=500, detail=f"PDF 변환 중 오류가 발생했습니다: {str(e)}")

@router.post("/{report_id}/regenerate-pdf")
async def regenerate_report_pdf(
    report_id: int,
    user_id: str = Query(..., description="사용자 ID")
):
    logger.info(f"[VSS] regenerate_report_pdf 호출: report_id={report_id}, user_id={user_id}")
    """보고서 PDF를 강제로 재생성"""
    try:
        if not user_id:
            raise HTTPException(status_code=400, detail="사용자 ID가 필요합니다.")
        
        # PDF 변환은 document service에서 처리합니다.
        verify_user_exists(user_id)
        
        row = get_report(report_id, user_id)
        if not row:
            raise HTTPException(status_code=404, detail="보고서를 찾을 수 없거나 권한이 없습니다.")
        report_id_db, title, _, _, _, _, _, _, _ = row
        
        # Word 파일 경로 찾기
        pattern = str(REPORTS_DIR / f"report_{report_id_db}_*.docx")
        matching_files = glob.glob(pattern)
        if not matching_files:
            raise HTTPException(status_code=404, detail="Word 파일을 찾을 수 없습니다.")
        
        docx_file_path = matching_files[0]
        pdf_file_path = str(Path(docx_file_path).with_suffix('.pdf'))
        
        # 기존 PDF 파일 삭제 (있는 경우)
        if os.path.exists(pdf_file_path):
            try:
                os.unlink(pdf_file_path)
            except Exception as e:
                logger.warning(f"기존 PDF 파일 삭제 실패: {e}")
        
        # PDF 변환 실행
        converted_pdf_path = DocumentService.convert_docx_to_pdf(docx_file_path, pdf_file_path)
        if not converted_pdf_path:
            raise HTTPException(status_code=500, detail="PDF 변환에 실패했습니다.")
        
        return {
            "success": True,
            "message": "PDF가 성공적으로 재생성되었습니다.",
            "pdf_url": f"/reports/{report_id}/pdf?user_id={user_id}"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"PDF 재생성 실패: {e}")
        raise HTTPException(status_code=500, detail=f"PDF 재생성 중 오류가 발생했습니다: {str(e)}")
